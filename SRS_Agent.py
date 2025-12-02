"""
Production-Ready SRS Generation Agent
Combines Groq Compound (web search) + Llama 3.3 70B (SRS generation) with document parsing
Author: SRS Generation Pipeline v1.0
"""

import os
import json
import time
import logging
from typing import Optional, Dict, List, Tuple
from datetime import datetime
from pathlib import Path
import re

# Third-party imports
from groq import Groq
import pypdf
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)


class DocumentParser:
    """Handles extraction from PDF and Word documents"""
    
    @staticmethod
    def parse_pdf(file_path: str) -> str:
        """Extract text from PDF file"""
        try:
            text = ""
            with open(file_path, 'rb') as file:
                pdf_reader = pypdf.PdfReader(file)
                logger.info(f"Extracting from PDF with {len(pdf_reader.pages)} pages")
                
                for page_num, page in enumerate(pdf_reader.pages):
                    text += f"\n--- Page {page_num + 1} ---\n"
                    text += page.extract_text()
            
            return text.strip()
        except Exception as e:
            logger.error(f"Error parsing PDF: {str(e)}")
            raise
    
    @staticmethod
    def parse_word(file_path: str) -> str:
        """Extract text from Word document"""
        try:
            doc = Document(file_path)
            text = ""
            
            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text += para.text + "\n"
            
            # Extract tables
            for table in doc.tables:
                text += "\n[TABLE]\n"
                for row in table.rows:
                    row_text = " | ".join(cell.text for cell in row.cells)
                    text += row_text + "\n"
            
            logger.info(f"Extracted {len(text)} characters from Word document")
            return text.strip()
        except Exception as e:
            logger.error(f"Error parsing Word document: {str(e)}")
            raise


class GroqSearchAgent:
    """Handles intelligent web searches via Groq Compound"""
    
    def __init__(self, api_key: Optional[str] = None):
        self.client = Groq(api_key=api_key or os.environ.get("GROQ_API_KEY"))
        self.model = "groq/compound"  # Compound model for web search
        self.rate_limit_delay = 0.1  # Small delay between requests
    
    def search_similar_projects(self, project_title: str, description: str) -> str:
        """Search for similar projects and case studies"""
        query = f"""
        Find information about projects similar to: {project_title}
        
        Project Description: {description[:500]}
        
        Please provide:
        1. 2-3 similar existing projects/products
        2. Their key features and technical approach
        3. Technologies they use
        4. Market positioning
        
        Keep response concise and factual.
        """
        
        try:
            logger.info("Searching for similar projects...")
            time.sleep(self.rate_limit_delay)
            
            response = self.client.chat.completions.create(
                model=self.model,
                messages=[{"role": "user", "content": query}],
                max_tokens=1500,
                temperature=0.3  # Lower temperature for factual research
            )
            
            return response.choices[0].message.content
        except Exception as e:
            logger.error(f"Search error: {str(e)}")
            return "Unable to retrieve similar projects information."
    
    def search_technology_stack(self, project_type: str, requirements: str) -> str:
        """Search for recommended technology stacks"""
        query = f"""
        Recommend technology stack for a {project_type} project.
        
        Requirements: {requirements[:500]}
        
        Provide:
        1. Recommended backend technologies
        2. Frontend technologies if applicable
        3. Databases and storage solutions
        4. DevOps and deployment tools
        5. Security and testing tools
        6. Justification for each recommendation
        
        Focus on industry best practices and proven solutions.
        """
        
        try:
            logger.info("Searching for technology recommendations...")
            time.sleep(self.rate_limit_delay)
            
            response = self.client.chat.completions.create(
                model=self.model,
                messages=[{"role": "user", "content": query}],
                max_tokens=2000,
                temperature=0.3
            )
            
            return response.choices[0].message.content
        except Exception as e:
            logger.error(f"Technology search error: {str(e)}")
            return "Unable to retrieve technology recommendations."
    
    def search_industry_benchmarks(self, industry: str, project_scope: str) -> str:
        """Search for industry benchmarks and cost data"""
        query = f"""
        Provide industry benchmarks for {industry} projects.
        
        Project Scope: {project_scope[:300]}
        
        Find and provide:
        1. Average project duration (weeks/months)
        2. Team size requirements
        3. Typical cost ranges
        4. Development effort estimation methods
        5. Risk factors commonly encountered
        
        Use recent data if available.
        """
        
        try:
            logger.info("Searching for industry benchmarks...")
            time.sleep(self.rate_limit_delay)
            
            response = self.client.chat.completions.create(
                model=self.model,
                messages=[{"role": "user", "content": query}],
                max_tokens=1500,
                temperature=0.3
            )
            
            return response.choices[0].message.content
        except Exception as e:
            logger.error(f"Benchmark search error: {str(e)}")
            return "Unable to retrieve industry benchmarks."


class SRSGenerator:
    """Generates comprehensive SRS using Llama 3.3 70B with reasoning"""
    
    def __init__(self, api_key: Optional[str] = None):
        self.client = Groq(api_key=api_key or os.environ.get("GROQ_API_KEY"))
        self.model = "llama-3.3-70b-versatile"
        self.rate_limit_delay = 0.2  # Slightly higher delay for stability
    
    def generate_srs(
        self,
        project_title: str,
        detailed_idea: str,
        document_content: Optional[str] = None,
        research_summary: Optional[str] = None
    ) -> Dict[str, str]:
        """
        Generate comprehensive SRS with extended thinking
        Returns dictionary with SRS sections
        """
        
        # Construct context for LLM
        context = f"""
PROJECT TITLE: {project_title}

DETAILED IDEA:
{detailed_idea}
"""
        
        if document_content:
            context += f"""
UPLOADED DOCUMENT CONTENT:
{document_content[:3000]}
"""
        
        if research_summary:
            context += f"""
RESEARCH FINDINGS (Similar Projects, Tech Stack, Benchmarks):
{research_summary}
"""
        
        # Detailed prompt for SRS generation with reasoning
        srs_prompt = f"""
You are an expert Software Requirements Specification writer with deep domain expertise.
Using the provided project details, research, and your knowledge, generate a COMPREHENSIVE 
Software Requirements Specification document.

{context}

Generate a detailed SRS document with the following sections:

1. **EXECUTIVE SUMMARY** (200-300 words)
   - Project overview
   - Key objectives
   - Business value
   - Success criteria

2. **PROJECT DESCRIPTION** (300-400 words)
   - Problem statement
   - Proposed solution
   - Target users/stakeholders
   - Project scope

3. **FUNCTIONAL REQUIREMENTS** (500-700 words)
   - List 8-12 functional requirements
   - Each with ID, title, description, priority (High/Medium/Low)
   - Format: [FR-001] Clear, measurable requirements

4. **NON-FUNCTIONAL REQUIREMENTS** (300-400 words)
   - Performance requirements
   - Security requirements
   - Scalability requirements
   - Reliability/Availability requirements
   - Maintainability requirements
   - Compliance requirements (if applicable)

5. **USE CASES** (400-500 words)
   - 4-6 primary use cases
   - Actor, preconditions, main flow, alternative flows, postconditions
   - Format each clearly

6. **USER STORIES** (300-400 words)
   - 8-10 user stories
   - Format: "As a [actor], I want to [action], so that [benefit]"
   - Include acceptance criteria for each

7. **DATA REQUIREMENTS** (200-300 words)
   - Key data entities
   - Data relationships
   - Data volume estimates
   - Data retention policies

8. **INTERFACE REQUIREMENTS** (200-300 words)
   - User interface expectations
   - External system integrations
   - API requirements
   - Third-party services

9. **ASSUMPTIONS & CONSTRAINTS** (200-300 words)
   - Technical assumptions
   - Business constraints
   - Resource constraints
   - Timeline assumptions

10. **ACCEPTANCE CRITERIA** (150-200 words)
    - Project completion criteria
    - Quality metrics
    - Testing requirements

---

CRITICAL INSTRUCTIONS FOR RESPONSE:
- Use clear, professional language
- Be specific and measurable, not vague
- Consider industry best practices
- Incorporate research findings naturally
- Ensure consistency across sections
- Each section should be well-developed and thorough
- Use JSON-compatible formatting (escaped quotes for embedded quotes)

Generate the complete SRS now:
"""
        
        try:
            logger.info("Generating SRS with Llama 3.3 70B (extended thinking)...")
            time.sleep(self.rate_limit_delay)
            
            response = self.client.chat.completions.create(
                model=self.model,
                messages=[{"role": "user", "content": srs_prompt}],
                max_tokens=8000,  # Large context for comprehensive SRS
                temperature=0.7,  # Balanced for quality and creativity
                top_p=0.9
            )
            
            srs_content = response.choices[0].message.content
            logger.info("SRS generation completed successfully")
            
            # Parse and structure the SRS content
            return self._structure_srs(srs_content)
        
        except Exception as e:
            logger.error(f"SRS generation error: {str(e)}")
            raise
    
    def _structure_srs(self, srs_text: str) -> Dict[str, str]:
        """Structure raw SRS text into sections"""
        sections = {
            "executive_summary": "",
            "project_description": "",
            "functional_requirements": "",
            "non_functional_requirements": "",
            "use_cases": "",
            "user_stories": "",
            "data_requirements": "",
            "interface_requirements": "",
            "assumptions_constraints": "",
            "acceptance_criteria": "",
            "raw_content": srs_text
        }
        
        # Simple section parsing (can be improved with regex)
        section_patterns = {
            "executive_summary": r"(?:1\.|EXECUTIVE SUMMARY.*?)(?=\n2\.|NON-FUNCTIONAL)",
            "project_description": r"(?:2\.|PROJECT DESCRIPTION.*?)(?=\n3\.|FUNCTIONAL)",
            "functional_requirements": r"(?:3\.|FUNCTIONAL REQUIREMENTS.*?)(?=\n4\.|NON-FUNCTIONAL)",
            "non_functional_requirements": r"(?:4\.|NON-FUNCTIONAL.*?)(?=\n5\.|USE CASES)",
            "use_cases": r"(?:5\.|USE CASES.*?)(?=\n6\.|USER STORIES)",
            "user_stories": r"(?:6\.|USER STORIES.*?)(?=\n7\.|DATA)",
            "data_requirements": r"(?:7\.|DATA.*?)(?=\n8\.|INTERFACE)",
            "interface_requirements": r"(?:8\.|INTERFACE.*?)(?=\n9\.|ASSUMPTIONS)",
            "assumptions_constraints": r"(?:9\.|ASSUMPTIONS.*?)(?=\n10\.|ACCEPTANCE)",
            "acceptance_criteria": r"(?:10\.|ACCEPTANCE.*?)$"
        }
        
        for section, pattern in section_patterns.items():
            match = re.search(pattern, srs_text, re.DOTALL | re.IGNORECASE)
            if match:
                sections[section] = match.group(0).strip()
        
        return sections


class WordDocumentGenerator:
    """Generates professional Word document from SRS data"""
    
    HEADING_STYLE_MAP = {
        1: "Heading 1",
        2: "Heading 2",
        3: "Heading 3"
    }
    
    def __init__(self):
        self.doc = Document()
        self.setup_styles()
    
    def setup_styles(self):
        """Configure document styles"""
        # Set default font
        style = self.doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)
    
    def add_title_page(self, project_title: str, timestamp: str):
        """Add professional title page"""
        title = self.doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.add_run(project_title)
        title_run.font.size = Pt(28)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(0, 51, 102)
        
        self.doc.add_paragraph()
        
        subtitle = self.doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_run = subtitle.add_run("Software Requirements Specification (SRS)")
        subtitle_run.font.size = Pt(16)
        subtitle_run.font.color.rgb = RGBColor(64, 64, 64)
        
        self.doc.add_paragraph()
        self.doc.add_paragraph()
        
        info = self.doc.add_paragraph()
        info.alignment = WD_ALIGN_PARAGRAPH.CENTER
        info_run = info.add_run(f"Generated on: {timestamp}\nVersion 1.0")
        info_run.font.size = Pt(12)
        info_run.font.italic = True
        
        self.doc.add_page_break()
    
    def add_table_of_contents(self):
        """Add table of contents"""
        self.doc.add_heading('Table of Contents', level=1)
        toc_items = [
            "Executive Summary",
            "Project Description",
            "Functional Requirements",
            "Non-Functional Requirements",
            "Use Cases",
            "User Stories",
            "Data Requirements",
            "Interface Requirements",
            "Assumptions & Constraints",
            "Acceptance Criteria"
        ]
        
        for item in toc_items:
            p = self.doc.add_paragraph(item, style='List Bullet')
            p.paragraph_format.left_indent = Inches(0.25)
        
        self.doc.add_page_break()
    
    def add_section(self, heading: str, content: str, level: int = 1):
        """Add section with heading and content"""
        self.doc.add_heading(heading, level=level)
        
        # Parse content for better formatting
        paragraphs = content.split('\n\n')
        for para_text in paragraphs:
            if para_text.strip():
                # Check if it's a requirement or list item
                if para_text.strip().startswith(('[', '•', '-', '*')):
                    para = self.doc.add_paragraph(para_text.strip(), style='List Bullet')
                else:
                    para = self.doc.add_paragraph(para_text.strip())
                
                para.paragraph_format.space_after = Pt(6)
    
    def add_requirements_table(self, requirements: str):
        """Add formatted requirements table"""
        # Parse requirements
        req_pattern = r'\[([A-Z]+-\d+)\]\s+(.+?)(?=\[|$)'
        matches = re.findall(req_pattern, requirements, re.DOTALL)
        
        if matches:
            table = self.doc.add_table(rows=1, cols=4)
            table.style = 'Light Grid Accent 1'
            
            # Header row
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'ID'
            hdr_cells[1].text = 'Description'
            hdr_cells[2].text = 'Priority'
            hdr_cells[3].text = 'Status'
            
            # Add requirement rows
            for req_id, req_text in matches:
                row_cells = table.add_row().cells
                row_cells[0].text = req_id
                row_cells[1].text = req_text.strip()[:100]
                row_cells[2].text = 'TBD'
                row_cells[3].text = 'Not Started'
    
    def add_footer(self):
        """Add document footer"""
        section = self.doc.sections[0]
        footer = section.footer
        footer_para = footer.paragraphs[0]
        footer_para.text = f"Confidential - Software Requirements Specification | Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    def generate(self, srs_data: Dict[str, str], project_title: str, output_path: str):
        """Generate complete Word document"""
        timestamp = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        
        # Add title page
        self.add_title_page(project_title, timestamp)
        
        # Add table of contents
        self.add_table_of_contents()
        
        # Add sections
        sections = [
            ("Executive Summary", srs_data.get("executive_summary", ""), 1),
            ("Project Description", srs_data.get("project_description", ""), 1),
            ("Functional Requirements", srs_data.get("functional_requirements", ""), 1),
            ("Non-Functional Requirements", srs_data.get("non_functional_requirements", ""), 1),
            ("Use Cases", srs_data.get("use_cases", ""), 1),
            ("User Stories", srs_data.get("user_stories", ""), 1),
            ("Data Requirements", srs_data.get("data_requirements", ""), 1),
            ("Interface Requirements", srs_data.get("interface_requirements", ""), 1),
            ("Assumptions & Constraints", srs_data.get("assumptions_constraints", ""), 1),
            ("Acceptance Criteria", srs_data.get("acceptance_criteria", ""), 1),
        ]
        
        for heading, content, level in sections:
            if content:
                self.add_section(heading, content, level)
                
                # Add page break between major sections
                if level == 1 and heading != sections[-1][0]:
                    self.doc.add_page_break()
        
        # Add footer
        self.add_footer()
        
        # Save document
        self.doc.save(output_path)
        logger.info(f"SRS document saved to: {output_path}")


class SRSPipeline:
    """Main orchestration pipeline"""
    
    def __init__(self):
        self.doc_parser = DocumentParser()
        self.search_agent = GroqSearchAgent()
        self.srs_generator = SRSGenerator()
    
    def process(
        self,
        project_title: str,
        detailed_idea: str,
        attachment_path: Optional[str] = None,
        output_dir: str = "./output"
    ) -> str:
        """
        Main processing pipeline
        
        Args:
            project_title: Name of the project
            detailed_idea: Detailed description from user
            attachment_path: Path to PDF or Word document (optional)
            output_dir: Output directory for generated SRS
        
        Returns:
            Path to generated SRS document
        """
        
        logger.info(f"Starting SRS generation for: {project_title}")
        
        # Create output directory
        Path(output_dir).mkdir(parents=True, exist_ok=True)
        
        # Step 1: Parse attachment if available
        document_content = None
        if attachment_path:
            logger.info(f"Parsing attachment: {attachment_path}")
            if attachment_path.lower().endswith('.pdf'):
                document_content = self.doc_parser.parse_pdf(attachment_path)
            elif attachment_path.lower().endswith('.docx'):
                document_content = self.doc_parser.parse_word(attachment_path)
            logger.info(f"Extracted {len(document_content)} characters from document")
        
        # Step 2: Conduct targeted web searches (minimize API calls)
        logger.info("Conducting targeted research...")
        research_summary = self._conduct_research(project_title, detailed_idea)
        
        # Step 3: Generate SRS
        logger.info("Generating SRS with Llama 3.3 70B...")
        srs_data = self.srs_generator.generate_srs(
            project_title=project_title,
            detailed_idea=detailed_idea,
            document_content=document_content,
            research_summary=research_summary
        )
        
        # Step 4: Generate Word document
        logger.info("Creating Word document...")
        output_file = os.path.join(
            output_dir,
            f"SRS_{project_title.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        )
        
        doc_gen = WordDocumentGenerator()
        doc_gen.generate(srs_data, project_title, output_file)
        
        logger.info(f"✓ SRS generation completed successfully!")
        logger.info(f"✓ Output file: {output_file}")
        
        return output_file
    
    def _conduct_research(self, project_title: str, detailed_idea: str) -> str:
        """Conduct focused research to minimize API calls"""
        
        research_parts = []
        
        # Search 1: Similar projects (most valuable)
        similar_projects = self.search_agent.search_similar_projects(
            project_title, detailed_idea
        )
        research_parts.append(f"## Similar Projects and Case Studies\n{similar_projects}")
        
        # Search 2: Technology stack
        tech_stack = self.search_agent.search_technology_stack(
            "software application", detailed_idea
        )
        research_parts.append(f"\n## Recommended Technology Stack\n{tech_stack}")
        
        # Search 3: Industry benchmarks (if project seems substantial)
        if len(detailed_idea) > 500:  # Only for detailed projects
            benchmarks = self.search_agent.search_industry_benchmarks(
                "software development", detailed_idea[:300]
            )
            research_parts.append(f"\n## Industry Benchmarks\n{benchmarks}")
        
        return "\n".join(research_parts)


# ============================================================================
# USAGE EXAMPLE
# ============================================================================

def main():
    """Example usage of the SRS Pipeline"""
    
    # Example 1: Without attachment
    pipeline = SRSPipeline()
    
    project_title = "AI-Powered Document Processing Platform"
    detailed_idea = """
    We need to build a platform that automatically processes scanned documents and PDFs 
    using AI/ML models to extract structured data, generate insights, and enable intelligent search.
    
    Key features:
    - Upload and process multiple document formats (PDF, images, scanned docs)
    - Extract key information using vision models and LLMs
    - Generate summaries and insights automatically
    - Full-text search with semantic understanding
    - User-friendly dashboard for document management
    - API for third-party integrations
    - Export results in multiple formats (JSON, CSV, PDF)
    
    Target users: Enterprise customers, legal firms, healthcare providers
    """
    
    # Generate SRS
    output_file = pipeline.process(
        project_title=project_title,
        detailed_idea=detailed_idea,
        attachment_path=None,  # Optional: provide path to PDF/DOCX
        output_dir="./srs_output"
    )
    
    print(f"\n✓ SRS Document generated: {output_file}")


if __name__ == "__main__":
    main()