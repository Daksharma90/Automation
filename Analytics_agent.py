"""
Production-Ready Feasibility Analysis Agent
Analyzes SRS documents to generate feasibility reports, risk assessments, and project plans
Uses Groq Compound (web search) + Llama 3.3 70B (analysis & reporting)
Author: Feasibility Analysis Pipeline v1.0
"""

import os
import json
import time
import logging
from typing import Optional, Dict, List, Tuple
from datetime import datetime
from pathlib import Path
import re

from groq import Groq
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)


class SRSParser:
    """Extract and structure information from SRS documents"""
    
    @staticmethod
    def parse_srs_document(file_path: str) -> Dict[str, str]:
        """Parse SRS Word document and extract sections"""
        try:
            doc = Document(file_path)
            sections = {}
            current_section = None
            current_content = []
            
            for para in doc.paragraphs:
                text = para.text.strip()
                
                # Detect section headings (typically Heading 1 style)
                if para.style.name.startswith('Heading 1'):
                    # Save previous section
                    if current_section:
                        sections[current_section] = '\n'.join(current_content)
                    
                    current_section = text
                    current_content = []
                elif text and current_section:
                    current_content.append(text)
            
            # Save last section
            if current_section:
                sections[current_section] = '\n'.join(current_content)
            
            logger.info(f"Parsed SRS document with {len(sections)} sections")
            return sections
        
        except Exception as e:
            logger.error(f"Error parsing SRS document: {str(e)}")
            raise
    
    @staticmethod
    def extract_key_requirements(sections: Dict[str, str]) -> Dict[str, List[str]]:
        """Extract key requirements and constraints from SRS sections"""
        extracted = {
            "functional_requirements": [],
            "non_functional_requirements": [],
            "technologies": [],
            "constraints": [],
            "scope": ""
        }
        
        # Extract from functional requirements
        if "Functional Requirements" in sections:
            reqs = sections["Functional Requirements"]
            # Find requirement IDs (FR-001, etc.)
            req_ids = re.findall(r'\[FR-\d+\][^\n]*', reqs)
            extracted["functional_requirements"] = req_ids[:12]  # Top 12
        
        # Extract from non-functional requirements
        if "Non-Functional Requirements" in sections:
            nf_reqs = sections["Non-Functional Requirements"]
            extracted["non_functional_requirements"] = nf_reqs[:500]  # First 500 chars
        
        # Extract from project description
        if "Project Description" in sections:
            extracted["scope"] = sections["Project Description"][:1000]
        
        # Look for technology mentions
        all_text = ' '.join(sections.values())
        tech_keywords = ['Python', 'JavaScript', 'React', 'Node.js', 'PostgreSQL', 
                        'MongoDB', 'AWS', 'Azure', 'Docker', 'Kubernetes', 'microservices']
        extracted["technologies"] = [t for t in tech_keywords if t.lower() in all_text.lower()]
        
        return extracted


class FeasibilitySearchAgent:
    """Conducts targeted web searches for feasibility analysis"""
    
    def __init__(self, api_key: Optional[str] = None):
        self.client = Groq(api_key=api_key or os.environ.get("GROQ_API_KEY"))
        self.model = "groq/compound"
        self.rate_limit_delay = 0.1
    
    def search_technology_feasibility(self, technologies: List[str], requirements: str) -> str:
        """Search for technology feasibility and maturity"""
        tech_list = ", ".join(technologies[:5])  # Top 5 technologies
        
        query = f"""
        Assess the feasibility and maturity of these technologies: {tech_list}
        
        For the given requirements: {requirements[:300]}
        
        Provide:
        1. Maturity level of each technology (stable/emerging/experimental)
        2. Community support and ecosystem
        3. Known limitations or challenges
        4. Learning curve and developer availability
        5. Integration capabilities
        6. Production readiness
        
        Be specific and factual. Focus on real-world viability.
        """
        
        try:
            logger.info("Searching technology feasibility...")
            time.sleep(self.rate_limit_delay)
            
            response = self.client.chat.completions.create(
                model=self.model,
                messages=[{"role": "user", "content": query}],
                max_tokens=2000,
                temperature=0.3
            )
            
            return response.choices[0].message.content
        except Exception as e:
            logger.error(f"Technology feasibility search error: {str(e)}")
            return "Unable to retrieve technology feasibility information."
    
    def search_project_risks(self, project_type: str, scope: str, complexity: str) -> str:
        """Search for common risks in similar projects"""
        query = f"""
        Identify common risks and challenges in {project_type} projects.
        
        Project Scope: {scope[:300]}
        Complexity Level: {complexity}
        
        Provide:
        1. Technical risks (scalability, performance, integration)
        2. Resource risks (skill gaps, team size, dependencies)
        3. Timeline risks (scheduling, scope creep, dependencies)
        4. Operational risks (maintenance, support, documentation)
        5. Business risks (market, ROI, stakeholder alignment)
        6. Security and compliance risks
        
        For each risk, suggest:
        - Probability (High/Medium/Low)
        - Impact severity
        - Mitigation strategies
        """
        
        try:
            logger.info("Searching project risks...")
            time.sleep(self.rate_limit_delay)
            
            response = self.client.chat.completions.create(
                model=self.model,
                messages=[{"role": "user", "content": query}],
                max_tokens=2500,
                temperature=0.3
            )
            
            return response.choices[0].message.content
        except Exception as e:
            logger.error(f"Risk search error: {str(e)}")
            return "Unable to retrieve risk information."
    
    def search_effort_estimation(self, features_count: int, team_size: str) -> str:
        """Search for effort estimation benchmarks"""
        query = f"""
        Provide effort estimation benchmarks for a software project.
        
        Project Features: {features_count} functional requirements
        Expected Team Size: {team_size}
        
        Based on industry standards, provide:
        1. Average effort per feature (story points or hours)
        2. Development phases and typical distribution:
           - Design phase (% of total time)
           - Development phase (% of total time)
           - Testing phase (% of total time)
           - Deployment phase (% of total time)
        3. Contingency buffer recommendations
        4. Milestone estimation
        5. Resource allocation patterns
        6. Productivity factors affecting timeline
        
        Reference recent industry data if available.
        """
        
        try:
            logger.info("Searching effort estimation benchmarks...")
            time.sleep(self.rate_limit_delay)
            
            response = self.client.chat.completions.create(
                model=self.model,
                messages=[{"role": "user", "content": query}],
                max_tokens=2000,
                temperature=0.3
            )
            
            return response.choices[0].message.content
        except Exception as e:
            logger.error(f"Effort estimation search error: {str(e)}")
            return "Unable to retrieve effort estimation data."
    
    def search_cost_benchmarks(self, team_size: str, duration: str, location: str = "global") -> str:
        """Search for cost estimation benchmarks"""
        query = f"""
        Provide software development cost benchmarks.
        
        Parameters:
        - Team Size: {team_size}
        - Duration: {duration}
        - Location: {location}
        
        Provide:
        1. Average developer salary ranges by role (junior/mid/senior)
        2. Infrastructure costs (cloud, licenses, tools)
        3. Overhead costs (management, HR, facilities)
        4. Total cost of ownership estimates
        5. Cost per feature benchmarks
        6. Budget contingency recommendations (typically 15-25%)
        
        Include both onshore and offshore options if applicable.
        """
        
        try:
            logger.info("Searching cost benchmarks...")
            time.sleep(self.rate_limit_delay)
            
            response = self.client.chat.completions.create(
                model=self.model,
                messages=[{"role": "user", "content": query}],
                max_tokens=2000,
                temperature=0.3
            )
            
            return response.choices[0].message.content
        except Exception as e:
            logger.error(f"Cost benchmark search error: {str(e)}")
            return "Unable to retrieve cost benchmark data."


class FeasibilityAnalyzer:
    """Generates comprehensive feasibility analysis using Llama 3.3 70B"""
    
    def __init__(self, api_key: Optional[str] = None):
        self.client = Groq(api_key=api_key or os.environ.get("GROQ_API_KEY"))
        self.model = "llama-3.3-70b-versatile"
        self.rate_limit_delay = 0.2
    
    def generate_feasibility_analysis(
        self,
        srs_content: str,
        tech_feasibility: str,
        project_risks: str,
        effort_data: str,
        cost_data: str
    ) -> Dict[str, str]:
        """Generate comprehensive feasibility analysis"""
        
        analysis_prompt = f"""
You are an expert Software Feasibility Analyst and Project Manager with 15+ years of experience.
Analyze the following Software Requirements Specification (SRS) and research data to generate a comprehensive feasibility analysis report.

---
SRS SUMMARY:
{srs_content[:2000]}

---
RESEARCH DATA:

Technology Feasibility Assessment:
{tech_feasibility}

Project Risks Analysis:
{project_risks}

Effort Estimation Benchmarks:
{effort_data}

Cost Benchmarks:
{cost_data}

---

Generate a comprehensive feasibility analysis with the following sections:

1. **EXECUTIVE SUMMARY** (300-400 words)
   - Overall feasibility recommendation (Feasible / Feasible with Conditions / Not Feasible)
   - Feasibility Score (1-10 with justification)
   - Key success factors
   - Critical blockers (if any)
   - Recommendation and next steps

2. **TECHNICAL FEASIBILITY ASSESSMENT** (500-600 words)
   - Technology stack viability
   - Architecture feasibility
   - Performance and scalability considerations
   - Integration challenges and solutions
   - Technical debt and mitigation
   - Alternative approaches (if risks identified)
   - Technology recommendations
   Score: (1-10 with detailed breakdown)

3. **RESOURCE REQUIREMENTS** (400-500 words)
   - Team composition (roles and count)
     * Backend developers (estimated count)
     * Frontend developers (estimated count)
     * QA/Testing team
     * DevOps/Infrastructure
     * Project Manager
     * Architect/Tech Lead
   - Skill requirements and gaps
   - Key resource risks
   - Mitigation strategies
   - Outsourcing recommendations (if applicable)

4. **EFFORT AND TIMELINE ESTIMATION** (500-600 words)
   - Total development effort (in person-months or person-weeks)
   - Phase-wise breakdown:
     * Requirements & Design: X% (~Y weeks)
     * Development: X% (~Y weeks)
     * Testing & QA: X% (~Y weeks)
     * Deployment & Documentation: X% (~Y weeks)
   - Critical path analysis
   - Milestone schedule (monthly/quarterly)
   - Dependencies and blockers
   - Contingency recommendations (suggest 20-30% buffer)
   - Optimistic / Realistic / Pessimistic timeline scenarios

5. **COST ESTIMATION** (400-500 words)
   - Development costs breakdown:
     * Personnel costs (by role and level)
     * Infrastructure and tools (~$X/month)
     * Third-party services and licenses
     * Training and development
   - Total estimated cost (base estimate)
   - Cost scenarios (optimistic/realistic/pessimistic)
   - Cost per feature analysis
   - Funding recommendations
   - ROI analysis (if business metrics provided)
   - Cost control strategies

6. **RISK ANALYSIS & MITIGATION** (600-700 words)
   
   HIGH PRIORITY RISKS:
   For each high-risk item, provide:
   - Risk ID (RISK-001, etc.)
   - Description
   - Probability (High/Medium/Low)
   - Impact (Critical/Major/Moderate/Minor)
   - Risk Score (Probability × Impact)
   - Mitigation Strategy
   - Contingency Plan
   - Owner/Responsible Party
   
   Cover these areas:
   - Technical Risks (scalability, performance, integration issues)
   - Resource Risks (key person dependency, skill gaps, turnover)
   - Timeline Risks (unrealistic schedules, scope creep, dependencies)
   - Operational Risks (maintenance, support, monitoring)
   - Market/Business Risks (adoption, competition, market timing)
   - Compliance/Security Risks (data protection, regulatory)
   
   MEDIUM PRIORITY RISKS:
   - List briefly with key mitigation
   
   Risk Matrix:
   - Create visual representation (text-based) of risk prioritization
   - Likelihood vs Impact quadrant analysis

7. **ASSUMPTIONS & CONSTRAINTS** (300-400 words)
   - Key assumptions made in this analysis
   - Technical constraints and limitations
   - Resource constraints
   - Budget constraints
   - Timeline constraints
   - External dependencies
   - Regulatory/compliance constraints
   - Known limitations of this analysis

8. **RECOMMENDATIONS & NEXT STEPS** (300-400 words)
   - GO / NO-GO recommendation with conditions
   - Phase 1 recommendations (what to start with)
   - Risk mitigation priorities (first 30 days actions)
   - Success criteria for feasibility approval
   - Recommended approach (Agile/Waterfall/Hybrid)
   - Pre-project activities
   - Success metrics and KPIs
   - Decision criteria for proceeding

9. **FEASIBILITY SCORECARD** (200-300 words)
   Create detailed scoring across:
   - Technical Feasibility (1-10): _/10
   - Resource Availability (1-10): _/10
   - Timeline Realism (1-10): _/10
   - Cost Estimation Confidence (1-10): _/10
   - Risk Manageability (1-10): _/10
   - Overall Feasibility (1-10): _/10
   
   Provide brief justification for each score.

10. **APPENDIX: DETAILED ASSUMPTIONS** (200-300 words)
    - Development environment assumptions
    - Team productivity assumptions (points/week)
    - Tool and infrastructure assumptions
    - Build vs Buy vs Outsource assumptions
    - Support and maintenance assumptions

---

CRITICAL REQUIREMENTS FOR RESPONSE:
- Use professional, clear language
- Be specific with numbers and metrics (not vague)
- Ground recommendations in industry best practices
- Clearly distinguish between feasible and concerning areas
- Provide actionable next steps
- Include contingency and risk-aware planning
- Ensure all sections cross-reference supporting data
- Maintain consistency in recommendations across sections

Generate the complete feasibility analysis now:
"""
        
        try:
            logger.info("Generating feasibility analysis with Llama 3.3 70B...")
            time.sleep(self.rate_limit_delay)
            
            response = self.client.chat.completions.create(
                model=self.model,
                messages=[{"role": "user", "content": analysis_prompt}],
                max_tokens=10000,
                temperature=0.7,
                top_p=0.9
            )
            
            analysis_content = response.choices[0].message.content
            logger.info("Feasibility analysis generated successfully")
            
            return self._structure_analysis(analysis_content)
        
        except Exception as e:
            logger.error(f"Analysis generation error: {str(e)}")
            raise
    
    def _structure_analysis(self, analysis_text: str) -> Dict[str, str]:
        """Structure raw analysis text into sections"""
        sections = {
            "executive_summary": "",
            "technical_feasibility": "",
            "resource_requirements": "",
            "effort_timeline": "",
            "cost_estimation": "",
            "risk_analysis": "",
            "assumptions_constraints": "",
            "recommendations": "",
            "feasibility_scorecard": "",
            "appendix": "",
            "raw_content": analysis_text
        }
        
        # Parse sections (simple pattern matching)
        section_patterns = {
            "executive_summary": r"(?:1\.|EXECUTIVE SUMMARY.*?)(?=\n2\.|TECHNICAL)",
            "technical_feasibility": r"(?:2\.|TECHNICAL FEASIBILITY.*?)(?=\n3\.|RESOURCE)",
            "resource_requirements": r"(?:3\.|RESOURCE REQUIREMENTS.*?)(?=\n4\.|EFFORT)",
            "effort_timeline": r"(?:4\.|EFFORT.*?)(?=\n5\.|COST)",
            "cost_estimation": r"(?:5\.|COST ESTIMATION.*?)(?=\n6\.|RISK)",
            "risk_analysis": r"(?:6\.|RISK ANALYSIS.*?)(?=\n7\.|ASSUMPTIONS)",
            "assumptions_constraints": r"(?:7\.|ASSUMPTIONS.*?)(?=\n8\.|RECOMMENDATIONS)",
            "recommendations": r"(?:8\.|RECOMMENDATIONS.*?)(?=\n9\.|FEASIBILITY)",
            "feasibility_scorecard": r"(?:9\.|FEASIBILITY SCORECARD.*?)(?=\n10\.|APPENDIX)",
            "appendix": r"(?:10\.|APPENDIX.*?)$"
        }
        
        for section, pattern in section_patterns.items():
            match = re.search(pattern, analysis_text, re.DOTALL | re.IGNORECASE)
            if match:
                sections[section] = match.group(0).strip()
        
        return sections


class FeasibilityReportGenerator:
    """Generates professional Word document for feasibility report"""
    
    def __init__(self):
        self.doc = Document()
        self.setup_styles()
        self.setup_table_style()
    
    def setup_styles(self):
        """Configure document styles"""
        style = self.doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)
    
    def setup_table_style(self):
        """Setup table styling"""
        pass
    
    def add_title_page(self, srs_title: str, timestamp: str):
        """Add professional title page"""
        title = self.doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.add_run("FEASIBILITY ANALYSIS REPORT")
        title_run.font.size = Pt(28)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(0, 51, 102)
        
        self.doc.add_paragraph()
        
        project = self.doc.add_paragraph()
        project.alignment = WD_ALIGN_PARAGRAPH.CENTER
        project_run = project.add_run(f"Project: {srs_title}")
        project_run.font.size = Pt(18)
        project_run.font.bold = True
        
        self.doc.add_paragraph()
        
        info = self.doc.add_paragraph()
        info.alignment = WD_ALIGN_PARAGRAPH.CENTER
        info_run = info.add_run(f"Generated on: {timestamp}\nVersion 1.0\n\n")
        info_run.font.size = Pt(12)
        info_run.font.italic = True
        
        self.doc.add_page_break()
    
    def add_executive_summary_box(self, content: str):
        """Add executive summary in highlighted box"""
        self.doc.add_heading('Executive Summary', level=1)
        
        # Extract feasibility score if present
        score_match = re.search(r'Feasibility Score[:\s]+(\d+)/10', content)
        recommendation_match = re.search(r'recommendation[:\s]+(\w+(?:\s+\w+)*)', content, re.IGNORECASE)
        
        if score_match:
            score = int(score_match.group(1))
            # Add score box
            p = self.doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            
            run = p.add_run(f"FEASIBILITY SCORE: {score}/10  |  ")
            run.font.bold = True
            run.font.size = Pt(12)
            
            if score >= 7:
                color_run = p.add_run("✓ FEASIBLE")
                color_run.font.color.rgb = RGBColor(34, 139, 34)  # Green
            elif score >= 5:
                color_run = p.add_run("⚠ FEASIBLE WITH CONDITIONS")
                color_run.font.color.rgb = RGBColor(255, 165, 0)  # Orange
            else:
                color_run = p.add_run("✗ NOT FEASIBLE")
                color_run.font.color.rgb = RGBColor(220, 20, 60)  # Red
            
            color_run.font.bold = True
            color_run.font.size = Pt(12)
        
        self.doc.add_paragraph(content)
    
    def add_risk_matrix_table(self, risk_content: str):
        """Add risk matrix table"""
        self.doc.add_heading('Risk Priority Matrix', level=2)
        
        # Create 4x4 matrix
        table = self.doc.add_table(rows=5, cols=5)
        table.style = 'Light Grid Accent 1'
        
        # Header rows and columns
        table.rows[0].cells[0].text = 'Risk Level'
        table.rows[0].cells[1].text = 'Low Impact'
        table.rows[0].cells[2].text = 'Medium Impact'
        table.rows[0].cells[3].text = 'High Impact'
        table.rows[0].cells[4].text = 'Critical'
        
        table.rows[1].cells[0].text = 'Low Probability'
        table.rows[2].cells[0].text = 'Medium Probability'
        table.rows[3].cells[0].text = 'High Probability'
        table.rows[4].cells[0].text = 'Certain'
        
        # Color cells
        for i in range(1, 5):
            for j in range(1, 5):
                cell = table.rows[i].cells[j]
                if i + j > 6:
                    # Red zone (high risk)
                    shading_elm = OxmlElement('w:shd')
                    shading_elm.set(qn('w:fill'), 'FFE0E0')
                    cell._element.get_or_add_tcPr().append(shading_elm)
                elif i + j >= 5:
                    # Orange zone (medium risk)
                    shading_elm = OxmlElement('w:shd')
                    shading_elm.set(qn('w:fill'), 'FFFACD')
                    cell._element.get_or_add_tcPr().append(shading_elm)
    
    def add_section(self, heading: str, content: str, level: int = 1):
        """Add section with heading and content"""
        self.doc.add_heading(heading, level=level)
        
        paragraphs = content.split('\n\n')
        for para_text in paragraphs:
            if para_text.strip():
                # Check if it's a list item
                if para_text.strip().startswith(('[', '•', '-', '*', '✓', '✗')):
                    para = self.doc.add_paragraph(para_text.strip(), style='List Bullet')
                else:
                    para = self.doc.add_paragraph(para_text.strip())
                
                para.paragraph_format.space_after = Pt(6)
    
    def add_scorecard_table(self, scorecard_content: str):
        """Add feasibility scorecard as table"""
        self.doc.add_heading('Feasibility Scorecard', level=2)
        
        table = self.doc.add_table(rows=7, cols=3)
        table.style = 'Light Grid Accent 1'
        
        # Header
        table.rows[0].cells[0].text = 'Criteria'
        table.rows[0].cells[1].text = 'Score'
        table.rows[0].cells[2].text = 'Status'
        
        # Criteria and default scores
        criteria = [
            'Technical Feasibility',
            'Resource Availability',
            'Timeline Realism',
            'Cost Confidence',
            'Risk Manageability',
            'Overall Feasibility'
        ]
        
        for idx, criterion in enumerate(criteria, 1):
            table.rows[idx].cells[0].text = criterion
            # Extract score from content if possible
            score_match = re.search(rf'{criterion}[:\s]+(\d+)/10', scorecard_content, re.IGNORECASE)
            if score_match:
                score = int(score_match.group(1))
                table.rows[idx].cells[1].text = f"{score}/10"
                
                if score >= 7:
                    table.rows[idx].cells[2].text = "✓ Good"
                elif score >= 5:
                    table.rows[idx].cells[2].text = "⚠ Fair"
                else:
                    table.rows[idx].cells[2].text = "✗ Poor"
    
    def add_footer(self):
        """Add document footer"""
        section = self.doc.sections[0]
        footer = section.footer
        footer_para = footer.paragraphs[0]
        footer_para.text = f"Confidential - Feasibility Analysis Report | Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    def generate(self, analysis_data: Dict[str, str], project_title: str, output_path: str):
        """Generate complete feasibility report"""
        timestamp = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        
        # Add title page
        self.add_title_page(project_title, timestamp)
        
        # Add table of contents
        toc_items = [
            "Executive Summary",
            "Technical Feasibility Assessment",
            "Resource Requirements",
            "Effort and Timeline Estimation",
            "Cost Estimation",
            "Risk Analysis & Mitigation",
            "Assumptions & Constraints",
            "Recommendations & Next Steps",
            "Feasibility Scorecard"
        ]
        
        self.doc.add_heading('Table of Contents', level=1)
        for item in toc_items:
            p = self.doc.add_paragraph(item, style='List Bullet')
            p.paragraph_format.left_indent = Inches(0.25)
        
        self.doc.add_page_break()
        
        # Add sections
        sections_to_add = [
            ("Executive Summary", analysis_data.get("executive_summary", ""), 1, True),
            ("Technical Feasibility Assessment", analysis_data.get("technical_feasibility", ""), 1, False),
            ("Resource Requirements", analysis_data.get("resource_requirements", ""), 1, False),
            ("Effort and Timeline Estimation", analysis_data.get("effort_timeline", ""), 1, False),
            ("Cost Estimation", analysis_data.get("cost_estimation", ""), 1, False),
            ("Risk Analysis & Mitigation", analysis_data.get("risk_analysis", ""), 1, False),
            ("Assumptions & Constraints", analysis_data.get("assumptions_constraints", ""), 1, False),
            ("Recommendations & Next Steps", analysis_data.get("recommendations", ""), 1, False),
            ("Feasibility Scorecard", analysis_data.get("feasibility_scorecard", ""), 1, True),
        ]
        
        for heading, content, level, is_special in sections_to_add:
            if content:
                if is_special and "Executive" in heading:
                    self.add_executive_summary_box(content)
                elif is_special and "Scorecard" in heading:
                    self.add_scorecard_table(content)
                else:
                    self.add_section(heading, content, level)
                
                # Add page break between major sections
                if level == 1 and heading != sections_to_add[-1][0]:
                    self.doc.add_page_break()
        
        # Add footer
        self.add_footer()
        
        # Save document
        self.doc.save(output_path)
        logger.info(f"Feasibility report saved to: {output_path}")


class FeasibilityPipeline:
    """Main orchestration pipeline for feasibility analysis"""
    
    def __init__(self):
        self.srs_parser = SRSParser()
        self.search_agent = FeasibilitySearchAgent()
        self.analyzer = FeasibilityAnalyzer()
    
    def process(
        self,
        srs_file_path: str,
        output_dir: str = "./output",
        project_title: Optional[str] = None
    ) -> str:
        """
        Main processing pipeline for feasibility analysis
        
        Args:
            srs_file_path: Path to generated SRS Word document
            output_dir: Directory for output feasibility report
            project_title: Optional project title (extracted from SRS if not provided)
        
        Returns:
            Path to generated feasibility report
        """
        
        logger.info(f"Starting feasibility analysis for SRS: {srs_file_path}")
        
        # Create output directory
        Path(output_dir).mkdir(parents=True, exist_ok=True)
        
        # Step 1: Parse SRS document
        logger.info("Parsing SRS document...")
        srs_sections = self.srs_parser.parse_srs_document(srs_file_path)
        
        if not srs_sections:
            raise ValueError("Failed to extract sections from SRS document")
        
        # Step 2: Extract key information
        logger.info("Extracting key requirements and constraints...")
        extracted_info = self.srs_parser.extract_key_requirements(srs_sections)
        
        # Step 3: Conduct targeted research
        logger.info("Conducting feasibility research...")
        
        # Research 1: Technology feasibility
        tech_feasibility = self.search_agent.search_technology_feasibility(
            extracted_info["technologies"],
            extracted_info["scope"]
        )
        
        # Research 2: Project risks
        project_type = "software application"  # Can be extracted from SRS
        project_risks = self.search_agent.search_project_risks(
            project_type,
            extracted_info["scope"],
            "medium"  # Can be calculated
        )
        
        # Research 3: Effort estimation
        features_count = len(extracted_info["functional_requirements"])
        effort_data = self.search_agent.search_effort_estimation(
            features_count,
            "5-8 developers"  # Default, can be customized
        )
        
        # Research 4: Cost benchmarks
        cost_data = self.search_agent.search_cost_benchmarks(
            "5-8 developers",
            "6-12 months",
            "global"
        )
        
        # Step 4: Generate feasibility analysis
        logger.info("Generating feasibility analysis...")
        srs_summary = self._create_srs_summary(srs_sections)
        
        analysis_data = self.analyzer.generate_feasibility_analysis(
            srs_content=srs_summary,
            tech_feasibility=tech_feasibility,
            project_risks=project_risks,
            effort_data=effort_data,
            cost_data=cost_data
        )
        
        # Step 5: Generate Word document
        logger.info("Creating feasibility report document...")
        if not project_title:
            project_title = list(srs_sections.keys())[0] if srs_sections else "Project"
        
        output_file = os.path.join(
            output_dir,
            f"FeasibilityReport_{project_title.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        )
        
        report_gen = FeasibilityReportGenerator()
        report_gen.generate(analysis_data, project_title, output_file)
        
        logger.info(f"✓ Feasibility analysis completed successfully!")
        logger.info(f"✓ Output file: {output_file}")
        
        return output_file
    
    def _create_srs_summary(self, srs_sections: Dict[str, str]) -> str:
        """Create compact SRS summary for analysis"""
        summary_parts = []
        
        key_sections = [
            "Executive Summary",
            "Project Description",
            "Functional Requirements",
            "Non-Functional Requirements"
        ]
        
        for section in key_sections:
            if section in srs_sections:
                content = srs_sections[section][:500]  # Limit to 500 chars
                summary_parts.append(f"## {section}\n{content}")
        
        return "\n\n".join(summary_parts)


# ============================================================================
# USAGE EXAMPLE
# ============================================================================

def main():
    """Example usage of the Feasibility Pipeline"""
    
    pipeline = FeasibilityPipeline()
    
    # Assume we have an SRS document from the previous phase
    srs_file = "./output/SRS_AI_Powered_Document_Processing_Platform_20251201_120000.docx"
    
    if not os.path.exists(srs_file):
        logger.warning(f"SRS file not found: {srs_file}")
        logger.info("To use this agent:")
        logger.info("1. Generate an SRS using srs_agent.py")
        logger.info("2. Update srs_file path above")
        logger.info("3. Run this script")
        return
    
    # Generate feasibility report
    output_file = pipeline.process(
        srs_file_path=srs_file,
        output_dir="./feasibility_output",
        project_title="AI Document Processing Platform"
    )
    
    print(f"\n✓ Feasibility Report generated: {output_file}")


if __name__ == "__main__":
    main()